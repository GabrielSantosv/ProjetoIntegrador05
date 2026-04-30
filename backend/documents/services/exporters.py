"""Excel and Word export builders."""
from io import BytesIO

from docx import Document as WordDocument
from openpyxl import Workbook
from openpyxl.styles import Font

from documents.models import Document


def build_excel(document: Document) -> BytesIO:
    """Build an .xlsx stream with extracted data and entities."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Dados extraidos"
    ws.append(["Campo", "Valor"])
    ws["A1"].font = ws["B1"].font = Font(bold=True)

    rows = {
        "Titulo": document.title,
        "Status": document.status,
        "Tipo": document.document_type,
        "Risco": document.risk_score,
        **document.extracted_data,
    }
    for key, value in rows.items():
        if key != "pages":
            ws.append([key, str(value)])

    entity_sheet = wb.create_sheet("Entidades")
    entity_sheet.append(["Rotulo", "Texto", "Confianca"])
    for entity in document.entities:
        entity_sheet.append([entity.get("label"), entity.get("text"), entity.get("score")])

    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream


def build_word(document: Document) -> BytesIO:
    """Build a .docx stream containing a concise legal processing report."""
    doc = WordDocument()
    doc.add_heading("Relatorio de Processamento Juridico", level=1)
    doc.add_paragraph(f"Documento: {document.title}")
    doc.add_paragraph(f"Tipo: {document.document_type or 'Nao classificado'}")
    doc.add_paragraph(f"Risco: {document.risk_score}/100")

    doc.add_heading("Dados extraidos", level=2)
    for key, value in document.extracted_data.items():
        if key != "pages":
            doc.add_paragraph(f"{key}: {value}")

    doc.add_heading("Parecer juridico IA", level=2)
    doc.add_paragraph(document.legal_opinion or "Sem parecer gerado.")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
